# Astra Resume Engine v3.2 (Final Production)
# Generation: gemini-3-flash-preview | ATS Score: gemini-3.1-flash-lite-preview
# API Keys: streamlit secrets | No Groq dependency
import streamlit as st
import json
import re
import io
import ast
import datetime
from typing import List

from pydantic import BaseModel, Field
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from xml.sax.saxutils import escape

# ============================================================
# 1. CONFIGURATION
# ============================================================
PAGE_TITLE = "Astra Resume Engine"
GEMINI_MODEL = "gemini-3-flash-preview"           # Main generation
SCORER_MODEL = "gemini-3.1-flash-lite-preview"     # ATS scoring (cheap + fast)

# Generic company names that mean "not found"
GENERIC_COMPANY_NAMES = {
    "", "company", "the company", "unknown", "n/a", "not specified",
    "not found", "undisclosed", "confidential", "hiring company",
    "saas company", "tech company", "startup", "enterprise",
    "entertainment company", "digital company", "global company",
}

ASTRA_PROMPT = """
Role: You are Astra, a Senior Technical Recruiter and Career Strategist.
Objective: Rewrite the resume to match the Job Description (JD) with 98% alignment, focusing on Narrative Fit, not just Keyword Matching.

CRITICAL STRATEGIC INSTRUCTIONS:

1. THE "DOMAIN BRIDGE" PROTOCOL:
   - Analyze the JD's specific industry context.
   - If the Candidate's experience is in a different domain, REFRAME experience to highlight TRANSFERABLE MECHANISMS.
   - Do not just list what they did; explicitly frame bullet points to solve the JD's specific business problem.

2. THE "PHILOSOPHY MATCH" PROTOCOL:
   - Analyze the JD for "Cultural Vibe" keywords.
   - If the JD asks for "Simple/Scrappy," DO NOT over-emphasize complex architectures.

3. EXECUTION GUIDELINES:
   - SUMMARY: Punchy. Explicitly mention the target role and how background solves JD's primary "Pain Point."
   - SKILLS: Create 6-7 dense categories. Prioritize JD-specific tools.
   - EXPERIENCE: Do NOT drop any roles. Include ALL experience entries. Rewrite bullets with metrics.
   - TARGET COMPANY: Extract the exact company name from the JD. If no company name is explicitly stated, return "Company" as the value.

4. BULLET POINT FORMATTING:
   - Do NOT prefix bullet points with "- ", "• ", or any marker. Return plain text starting with an action verb.

5. SKILL DENSITY (CRITICAL FOR ATS):
   - For each skill category, include the core tools from the JD AND closely related ecosystem tools the candidate knows.
   - Example: If JD mentions "AWS SageMaker", also include "Lambda, S3, EC2, CloudWatch" in that category.
   - Example: If JD mentions "Python", also include "FastAPI, Flask, Pandas, NumPy" in that category.
   - Example: If JD mentions "Docker", also include "Kubernetes, Helm, Docker Compose" in that category.
   - Example: If JD mentions "PyTorch", also include "TensorFlow, Keras, Scikit-learn" in that category.
   - ONLY add tools that are genuinely related and that the candidate actually has in their resume.
   - Goal: Maximum ATS keyword coverage without adding random/unrelated tools.

6. ANTI-HALLUCINATION RULES:
   - NEVER invent companies, roles, degrees, or dates not present in the input resume.
   - NEVER fabricate metrics. If the original has a metric, you may reframe it. If not, describe impact qualitatively.
   - ALL experience entries from the input resume MUST appear in the output. Do not merge or drop roles.
"""

COVER_LETTER_PROMPT = """
Role: You are the candidate writing a direct, high-impact email to a Hiring Manager.
Objective: Write a cover letter that sounds 100% HUMAN, authentic, and specific.

CRITICAL RULES:
1. BANNED PHRASES: NEVER use: "I am writing to express my interest," "I am excited to apply," "Please find my resume attached," "testament to," "underscores," "pivotal," "realm," "tapestry."
2. THE OPENING: Start with a "Hook" — an observation about the company's specific challenge.
3. THE "WAR STORY": Tell ONE specific story from experience that proves you can solve their problem.
4. TONE: Confident, conversational, peer-to-peer.

STRUCTURE:
1. Salutation: "Dear Hiring Team,"
2. The Hook: Connect to company's pain point.
3. The Bridge: "This challenge resonates with me because..."
4. The Evidence: The "War Story" with specific tools.
5. The Closing: Brief. End with "Thank you,"

Return ONLY the letter body text. No markdown.
"""

# ============================================================
# 2. PYDANTIC SCHEMAS
# ============================================================
class ExperienceItem(BaseModel):
    role_title: str = Field(description="The job title")
    company: str = Field(description="The company name")
    dates: str = Field(description="Employment dates")
    location: str = Field(description="City or Remote")
    responsibilities: List[str] = Field(description="6-8 bullet points. Plain text, no dash/bullet prefix.")
    achievements: List[str] = Field(description="2-3 quantified wins. Plain text, no dash/bullet prefix.")

class EducationItem(BaseModel):
    degree: str = Field(description="Degree name")
    college: str = Field(description="University name")

class SkillCategory(BaseModel):
    category: str = Field(description="Skill category name")
    technologies: str = Field(description="Comma-separated tools/skills with ecosystem density")

class ResumeSchema(BaseModel):
    candidate_name: str = Field(description="Full Name")
    candidate_title: str = Field(description="Professional Title matching JD")
    contact_info: str = Field(description="Phone | Email | Location")
    summary: str = Field(description="Professional summary tailored to JD")
    skills: List[SkillCategory] = Field(description="6-7 dense skill categories with ecosystem tools")
    experience: List[ExperienceItem] = Field(description="ALL roles from input resume")
    education: List[EducationItem] = Field(description="Educational background")
    target_company: str = Field(description="Company name from JD. If not explicitly named, return 'Company'.")

class ATSScoreSchema(BaseModel):
    score: int = Field(description="ATS match score from 0-100")
    reasoning: str = Field(description="One sentence explaining the score")

# ============================================================
# 3. HELPERS
# ============================================================
def strip_bullet_prefix(text):
    if not isinstance(text, str): return str(text)
    return re.sub(r'^[\s]*[-\u2013\u2014\u2022*]\s+', '', text)

def split_to_bullets(text_or_list):
    if isinstance(text_or_list, str):
        items = text_or_list.split('\n')
    elif isinstance(text_or_list, list):
        items = text_or_list
    else:
        items = [str(text_or_list)] if text_or_list else []
    return [strip_bullet_prefix(item) for item in items if item and item.strip()]

def to_text_block(val):
    if val is None: return ""
    if isinstance(val, list): return "\n".join([str(x) for x in val])
    return str(val)

def clean_skill_string(skill_str):
    if not isinstance(skill_str, str): return str(skill_str)
    if skill_str.strip().startswith("["):
        try:
            list_match = re.search(r"\[(.*?)\]", skill_str)
            if list_match:
                actual_list = ast.literal_eval(list_match.group(0))
                extra = skill_str[list_match.end():].strip().lstrip(",").strip()
                clean = ", ".join([str(s) for s in actual_list])
                if extra: clean += f", {extra}"
                return clean
        except: pass
    return skill_str

def is_generic_company(name):
    """Check if target_company is a generic placeholder."""
    if not name:
        return True
    return name.strip().lower() in GENERIC_COMPANY_NAMES

def get_first_name(full_name):
    """Extract first name from full name for filename."""
    if not full_name:
        return "Candidate"
    parts = full_name.strip().split()
    return parts[0] if parts else "Candidate"

def make_filename(data):
    """Build filename: Name_Company if real company, else FirstName_Resume."""
    c_name = data.get('candidate_name', 'Candidate').strip()
    target = data.get('target_company', '')

    safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', c_name.replace(' ', '_'))

    if is_generic_company(target):
        first = get_first_name(c_name)
        safe_first = re.sub(r'[^a-zA-Z0-9_-]', '_', first)
        return f"{safe_first}_Resume"
    else:
        safe_co = re.sub(r'[^a-zA-Z0-9_-]', '_', target.strip())
        return f"{safe_name}_{safe_co}"

def normalize_schema(data):
    if not isinstance(data, dict):
        return {"summary": str(data), "skills": {}, "experience": []}
    n = {}
    n['candidate_name'] = data.get('candidate_name', '')
    n['candidate_title'] = data.get('candidate_title', '')
    n['contact_info'] = str(data.get('contact_info', ''))
    n['summary'] = data.get('summary', '')

    raw_skills = data.get('skills', {})
    n['skills'] = {}
    if isinstance(raw_skills, dict):
        for k, v in raw_skills.items():
            n['skills'][k] = clean_skill_string(str(v))
    elif isinstance(raw_skills, list):
        for item in raw_skills:
            if isinstance(item, dict):
                cat = item.get('category', '')
                tech = item.get('technologies', '')
                if cat and tech: n['skills'][cat] = tech

    n['experience'] = []
    for role in data.get('experience', []):
        if isinstance(role, dict):
            n['experience'].append({
                'role_title': role.get('role_title', ''),
                'company': role.get('company', ''),
                'dates': role.get('dates', ''),
                'location': role.get('location', ''),
                'responsibilities': [strip_bullet_prefix(r) for r in role.get('responsibilities', [])],
                'achievements': split_to_bullets(role.get('achievements', []))
            })

    n['education'] = []
    for edu in data.get('education', []):
        if isinstance(edu, dict):
            n['education'].append({'degree': edu.get('degree', ''), 'college': edu.get('college', '')})
    n['target_company'] = data.get('target_company', 'Company')
    return n

# ============================================================
# 4. ATS SCORING (Gemini 3.1 Flash Lite)
# ============================================================
def calculate_ats_score(google_key, resume_json, jd_text):
    """Score resume-JD match using Gemini 3.1 Flash Lite (cheap + fast)."""
    client = genai.Client(api_key=google_key)
    try:
        response = client.models.generate_content(
            model=SCORER_MODEL,
            contents=f"""You are an ATS (Applicant Tracking System). Compare this Resume against the Job Description.
Score the match from 0-100 based on: keyword overlap, experience relevance, skill alignment, and seniority fit.
Be strict — a perfect resume gets 90-95, not 100.

RESUME:
{str(resume_json)[:3000]}

JOB DESCRIPTION:
{jd_text[:2500]}""",
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=ATSScoreSchema
            )
        )
        result = json.loads(response.text)
        return {
            "score": result.get("score", 0),
            "reasoning": result.get("reasoning", "")
        }
    except Exception as e:
        return {"score": 0, "reasoning": f"Scoring Error: {str(e)}"}

# ============================================================
# 5. GENERATION
# ============================================================
def analyze_and_generate(google_key, resume_text, jd_text):
    client = genai.Client(api_key=google_key)
    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=f"{ASTRA_PROMPT}\n\nRESUME:\n{resume_text}\n\nJD:\n{jd_text}",
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=ResumeSchema
            )
        )
        raw_data = json.loads(response.text)
        data = raw_data.model_dump() if hasattr(raw_data, 'model_dump') else raw_data
        data = normalize_schema(data)

        # ATS Score via Gemini 3.1 Flash Lite
        judge = calculate_ats_score(google_key, data, jd_text)
        data['ats_score'] = judge.get('score', 0)
        data['ats_reason'] = judge.get('reasoning', '')
        return data
    except Exception as e:
        return {"error": f"Generation Error: {str(e)}"}

def generate_cover_letter(google_key, resume_data, jd_text):
    client = genai.Client(api_key=google_key)
    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=f"{COVER_LETTER_PROMPT}\n\nRESUME DATA:\n{str(resume_data)}\n\nJOB DESCRIPTION:\n{jd_text}",
        )
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# ============================================================
# 6. DOCX RENDERERS
# ============================================================
def set_font(run, size, bold=False):
    run.font.name = 'Times New Roman'; run.font.size = Pt(size); run.bold = bold
    try: run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    except: pass

def create_doc(data):
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)
    for txt, sz, b in [(data.get('candidate_name',''),28,True),(data.get('candidate_title',''),14,True),(data.get('contact_info',''),12,True)]:
        p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        run = p.add_run(to_text_block(txt))
        if sz == 28: run.font.all_caps = True
        set_font(run, sz, b)
    def add_sec(title):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(title), 12, True)
    def add_body(txt, bullet=False):
        style = 'List Bullet' if bullet else 'Normal'
        p = doc.add_paragraph(style=style); p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0); set_font(p.add_run(to_text_block(txt)), 12)
    add_sec("Professional Profile"); add_body(data.get('summary', ''))
    add_sec("Key Skills/ Tools & Technologies")
    for k, v in data.get('skills', {}).items():
        p = doc.add_paragraph(style='List Bullet'); p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0); set_font(p.add_run(f"{k}: "), 12, True); set_font(p.add_run(to_text_block(v)), 12)
    add_sec("Professional Experience")
    for role in data.get('experience', []):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
        line = f"{role.get('role_title')} | {role.get('company')} | {role.get('location')} | {role.get('dates')}"
        set_font(p.add_run(to_text_block(line)), 12, True)
        for r in split_to_bullets(role.get('responsibilities', [])): add_body(r, bullet=True)
        achs = split_to_bullets(role.get('achievements', []))
        if achs:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run("Achievements:"), 12, True)
            for a in achs: add_body(a, bullet=True)
    add_sec("Education")
    for edu in data.get('education', []):
        add_body(f"{edu.get('degree','')}, {edu.get('college','')}", bullet=True)
    return doc

def create_cover_letter_doc(cl_text, data):
    doc = Document(); s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)
    def add_line(text, bold=False, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
        if not text: return
        p = doc.add_paragraph(); p.alignment = align; p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(str(text)); run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = bold
    add_line(data.get('candidate_name','').upper(), bold=True, space_after=0)
    contact = data.get('contact_info', '')
    if "|" in contact:
        for part in contact.split('|'): add_line(part.strip(), space_after=0)
    else: add_line(contact, space_after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_line(datetime.date.today().strftime("%B %d, %Y"), space_after=12)
    for para in cl_text.split('\n'):
        if para.strip(): add_line(para.strip(), space_after=12, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    return doc

# ============================================================
# 7. PDF RENDERER
# ============================================================
def create_pdf(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.5*inch, rightMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    sn = ParagraphStyle('AN', parent=styles['Normal'], fontName='Times-Roman', fontSize=12, leading=14, alignment=TA_JUSTIFY, spaceAfter=0)
    shn = ParagraphStyle('AHN', parent=styles['Normal'], fontName='Times-Bold', fontSize=28, leading=30, alignment=TA_CENTER, spaceAfter=0)
    sht = ParagraphStyle('AHT', parent=styles['Normal'], fontName='Times-Bold', fontSize=14, leading=16, alignment=TA_CENTER, spaceAfter=0)
    shc = ParagraphStyle('AHC', parent=styles['Normal'], fontName='Times-Bold', fontSize=12, leading=14, alignment=TA_CENTER, spaceAfter=6)
    ss = ParagraphStyle('AS', parent=styles['Normal'], fontName='Times-Bold', fontSize=12, leading=14, alignment=TA_LEFT, spaceBefore=12, spaceAfter=2)
    def clean(txt):
        if txt is None: return ""
        return escape(to_text_block(txt)).replace('\n', '<br/>')
    el = []
    el.append(Paragraph(clean(data.get('candidate_name','')), shn))
    el.append(Paragraph(clean(data.get('candidate_title','')), sht))
    el.append(Paragraph(clean(data.get('contact_info','')), shc))
    el.append(Paragraph("Professional Profile", ss))
    el.append(Paragraph(clean(data.get('summary','')), sn))
    el.append(Paragraph("Key Skills/ Tools & Technologies", ss))
    skill_items = []
    for k, v in data.get('skills', {}).items():
        skill_items.append(ListItem(Paragraph(f"<b>{clean(k)}:</b> {clean(v)}", sn), leftIndent=0))
    if skill_items: el.append(ListFlowable(skill_items, bulletType='bullet', start='\u2022', leftIndent=15))
    el.append(Paragraph("Professional Experience", ss))
    for role in data.get('experience', []):
        line = f"{role.get('role_title')} | {role.get('company')} | {role.get('location')} | {role.get('dates')}"
        el.append(Paragraph(f"<b>{clean(line)}</b>", sn)); el.append(Spacer(1,2))
        bullets = []
        for r in split_to_bullets(role.get('responsibilities', [])):
            bullets.append(ListItem(Paragraph(clean(r), sn), leftIndent=0))
        if bullets: el.append(ListFlowable(bullets, bulletType='bullet', start='\u2022', leftIndent=15))
        achs = split_to_bullets(role.get('achievements', []))
        if achs:
            el.append(Paragraph("<b>Achievements:</b>", sn))
            ab = [ListItem(Paragraph(clean(a), sn), leftIndent=0) for a in achs]
            if ab: el.append(ListFlowable(ab, bulletType='bullet', start='\u2022', leftIndent=25))
        el.append(Spacer(1,6))
    el.append(Paragraph("Education", ss))
    eb = [ListItem(Paragraph(clean(f"{e.get('degree','')}, {e.get('college','')}"), sn), leftIndent=0) for e in data.get('education',[])]
    if eb: el.append(ListFlowable(eb, bulletType='bullet', start='\u2022', leftIndent=15))
    doc.build(el); buffer.seek(0)
    return buffer.getvalue()

# ============================================================
# 8. UI
# ============================================================
st.set_page_config(page_title=PAGE_TITLE, layout="wide", page_icon="\U0001f680", initial_sidebar_state="expanded")
st.markdown("""<style>
#MainMenu {visibility: hidden;} footer {visibility: hidden;}
.block-container {padding-top: 1.5rem;}
div.stButton > button:first-child {border-radius: 6px; font-weight: 600;}
div[data-testid="stMetricValue"] {font-size: 1.8rem;}
</style>""", unsafe_allow_html=True)

# --- API Key from Streamlit Secrets (single key powers everything) ---
google_key = st.secrets.get("GOOGLE_API_KEY", "")

if 'data' not in st.session_state: st.session_state['data'] = None
if 'saved_base' not in st.session_state: st.session_state['saved_base'] = ""
if 'saved_jd' not in st.session_state: st.session_state['saved_jd'] = ""
if 'cover_letter' not in st.session_state: st.session_state['cover_letter'] = None

with st.sidebar:
    st.header("\u2699\ufe0f Configuration")
    if not google_key:
        st.error("GOOGLE_API_KEY missing from secrets.toml")
        st.code('# .streamlit/secrets.toml\nGOOGLE_API_KEY = "your-gemini-key"', language="toml")
    else:
        st.success("API Key loaded from secrets")
    st.divider()
    st.caption(f"Generation: `{GEMINI_MODEL}`")
    st.caption(f"ATS Scorer: `{SCORER_MODEL}`")
    st.divider()
    if st.button("\U0001f5d1\ufe0f Reset", use_container_width=True):
        st.session_state['data'] = None; st.session_state['saved_base'] = ""
        st.session_state['saved_jd'] = ""; st.session_state['cover_letter'] = None
        st.rerun()
    st.caption("Astra v3.2 Final")

if not st.session_state['data']:
    st.markdown(f"<h1 style='text-align:center;'>{PAGE_TITLE}</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center;color:#666;'>Paste Resume + JD \u2192 Tailored Resume + Cover Letter</p>", unsafe_allow_html=True)
    st.divider()
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("\U0001f4cb Base Resume")
        base = st.text_area("Resume", st.session_state['saved_base'], height=400, label_visibility="collapsed")
    with c2:
        st.subheader("\U0001f4bc Job Description")
        jd = st.text_area("JD", st.session_state['saved_jd'], height=400, label_visibility="collapsed")

    if st.button("\u2728 Architect My Application", type="primary", use_container_width=True):
        if google_key and base and jd:
            st.session_state['saved_base'] = base; st.session_state['saved_jd'] = jd
            with st.spinner("Generating tailored resume + scoring ATS match..."):
                data = analyze_and_generate(google_key, base, jd)
                if "error" in data: st.error(data['error'])
                else: st.session_state['data'] = data; st.rerun()
        else:
            st.warning("Need GOOGLE_API_KEY in secrets.toml + Resume + JD.")
else:
    data = st.session_state['data']
    target = data.get('target_company', '')
    display_target = target if not is_generic_company(target) else "Role-Specific Application"

    c1, c2, c3 = st.columns([1, 4, 1])
    with c2: st.markdown(f"## \U0001f3af Target: {display_target}")
    with c3:
        score = data.get('ats_score', 0)
        st.metric("ATS Match", f"{score}%")
        if data.get('ats_reason'):
            st.caption(data['ats_reason'])

    tab_edit, tab_export, tab_cover = st.tabs(["\U0001f4dd Editor", "\U0001f680 Export", "\u270d\ufe0f Cover Letter"])

    with tab_edit:
        with st.form("edit_form"):
            st.subheader("Candidate Details")
            c1, c2, c3 = st.columns(3)
            data['candidate_name'] = c1.text_input("Name", to_text_block(data.get('candidate_name')))
            data['candidate_title'] = c2.text_input("Title", to_text_block(data.get('candidate_title')))
            data['contact_info'] = c3.text_input("Contact", to_text_block(data.get('contact_info')))
            data['summary'] = st.text_area("Summary", to_text_block(data.get('summary')), height=120)
            st.subheader("Skills")
            skills = data.get('skills', {}); new_skills = {}; s_cols = st.columns(2)
            for i, (k, v) in enumerate(skills.items()):
                new_skills[k] = s_cols[i%2].text_area(k, to_text_block(v), key=f"sk_{i}", height=80).replace('\n', ', ')
            data['skills'] = new_skills
            st.subheader("Experience")
            for i, role in enumerate(data.get('experience', [])):
                with st.expander(f"{role.get('role_title','')} @ {role.get('company','')}"):
                    c1, c2 = st.columns(2)
                    role['role_title'] = c1.text_input("Role", to_text_block(role.get('role_title')), key=f"rt_{i}")
                    role['company'] = c2.text_input("Company", to_text_block(role.get('company')), key=f"co_{i}")
                    c3, c4 = st.columns(2)
                    role['dates'] = c3.text_input("Dates", to_text_block(role.get('dates')), key=f"dt_{i}")
                    role['location'] = c4.text_input("Location", to_text_block(role.get('location')), key=f"lo_{i}")
                    role['responsibilities'] = st.text_area("Responsibilities", to_text_block(role.get('responsibilities')), height=200, key=f"rp_{i}")
                    role['achievements'] = st.text_area("Achievements", to_text_block(role.get('achievements')), height=100, key=f"ac_{i}")
            st.subheader("Education")
            for i, edu in enumerate(data.get('education', [])):
                c1, c2 = st.columns(2)
                edu['degree'] = c1.text_input("Degree", to_text_block(edu.get('degree')), key=f"ed_{i}")
                edu['college'] = c2.text_input("College", to_text_block(edu.get('college')), key=f"ec_{i}")
            if st.form_submit_button("\U0001f4be Save", type="primary"):
                st.session_state['data'] = data; st.success("Saved!"); st.rerun()

    with tab_export:
        st.subheader("\U0001f4e5 Downloads")
        fname = make_filename(data)
        st.caption(f"Filename: `{fname}.docx` / `{fname}.pdf`")

        c1, c2 = st.columns(2)
        doc = create_doc(data); bio = io.BytesIO(); doc.save(bio)
        c1.download_button("\U0001f4c4 Word Doc", bio.getvalue(), f"{fname}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
        try:
            pdf_data = create_pdf(data)
            c2.download_button("\U0001f4d5 PDF", pdf_data, f"{fname}.pdf", "application/pdf", type="secondary", use_container_width=True)
        except Exception as e: c2.error(f"PDF Error: {e}")

    with tab_cover:
        st.subheader("\u270d\ufe0f Cover Letter")
        if st.button("\u2728 Draft Cover Letter", type="primary"):
            if google_key and st.session_state['saved_jd']:
                with st.spinner("Drafting..."):
                    st.session_state['cover_letter'] = generate_cover_letter(google_key, data, st.session_state['saved_jd'])
            else: st.warning("Need API key + JD")
        if st.session_state['cover_letter']:
            st.text_area("Preview", st.session_state['cover_letter'], height=400)
            cl_doc = create_cover_letter_doc(st.session_state['cover_letter'], data)
            bio_cl = io.BytesIO(); cl_doc.save(bio_cl)
            cl_fname = f"Cover_Letter_{fname}"
            st.download_button("\U0001f4c4 Download Cover Letter", bio_cl.getvalue(), f"{cl_fname}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

    st.divider()
    c3, c4 = st.columns(2)
    if c3.button("\u267b\ufe0f Re-Optimize", use_container_width=True):
        if st.session_state['saved_base'] and st.session_state['saved_jd']:
            with st.spinner("Re-generating..."):
                data = analyze_and_generate(google_key, st.session_state['saved_base'], st.session_state['saved_jd'])
                if "error" in data: st.error(data['error'])
                else: st.session_state['data'] = data; st.rerun()
    if c4.button("New Application", use_container_width=True):
        st.session_state['data'] = None; st.session_state['saved_jd'] = ""; st.session_state['cover_letter'] = None; st.rerun()
